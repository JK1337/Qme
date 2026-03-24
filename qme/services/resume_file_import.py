"""Extract plain text from resume files users obtain from LinkedIn (PDF / DOCX)."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

MAX_RESUME_FILE_BYTES = 5 * 1024 * 1024


def text_from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        t = t.strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts).strip()


def text_from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    lines: list[str] = []
    for p in doc.paragraphs:
        line = (p.text or "").strip()
        if line:
            lines.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                lines.append(row_text)
    return "\n".join(lines).strip()


def extract_resume_text(*, data: bytes, filename: str | None) -> str:
    if len(data) > MAX_RESUME_FILE_BYTES:
        raise ValueError(f"File too large (max {MAX_RESUME_FILE_BYTES // (1024 * 1024)} MB).")

    ext = Path(filename or "").suffix.lower()
    is_pdf_magic = data[:4] == b"%PDF"
    is_zip_magic = len(data) >= 4 and data[:2] == b"PK"

    if ext == ".pdf" or (ext == "" and is_pdf_magic):
        if not is_pdf_magic:
            raise ValueError("That file does not look like a valid PDF.")
        return text_from_pdf(data)

    if ext == ".docx" or (ext == "" and is_zip_magic):
        try:
            return text_from_docx(data)
        except Exception as exc:  # noqa: BLE001 — user upload; surface friendly message
            raise ValueError("Could not read that Word file — use a .docx from LinkedIn or save as PDF.") from exc

    raise ValueError("Please upload a .pdf or .docx file (LinkedIn → Save to PDF, or a downloaded resume).")
