"""Generate styled DOCX resume exports from structured sections."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor

from qme.services.resume_sections import ResumeSections, parse_resume_sections


def _style_defaults(doc: Document, style_name: str) -> None:
    normal = doc.styles["Normal"]
    font = normal.font
    if style_name == "executive":
        font.name = "Georgia"
        font.size = Pt(11)
    else:
        font.name = "Calibri"
        font.size = Pt(11)


def _heading_color(style_name: str) -> RGBColor:
    if style_name == "modern":
        return RGBColor(10, 102, 194)
    if style_name == "executive":
        return RGBColor(0, 65, 130)
    return RGBColor(25, 25, 25)


def _add_section(doc: Document, title: str, body: str, style_name: str) -> None:
    text = (body or "").strip()
    if not text:
        return
    heading = doc.add_paragraph()
    run = heading.add_run(title.upper() if style_name == "modern" else title)
    run.bold = True
    run.font.color.rgb = _heading_color(style_name)
    run.font.size = Pt(12 if style_name != "executive" else 11)

    for line in text.splitlines():
        row = line.strip()
        if not row:
            continue
        if row.startswith(("- ", "* ", "• ")):
            doc.add_paragraph(row[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(row)


def build_resume_docx(text: str, style_name: str = "classic") -> bytes:
    parsed: ResumeSections = parse_resume_sections(text or "")
    doc = Document()
    _style_defaults(doc, style_name)

    _add_section(doc, "Profile", parsed.profile, style_name)
    _add_section(doc, "Strengths", parsed.strengths, style_name)
    _add_section(doc, "Education", parsed.education, style_name)
    _add_section(doc, "Job Experience", parsed.job_experience, style_name)
    _add_section(doc, "Other", parsed.other, style_name)

    if len(doc.paragraphs) == 0:
        doc.add_paragraph("Resume content is empty.")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
